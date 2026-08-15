"""
Parse 7300+ GS English Medium DOCX - Fixed parser for tab-separated options
"""
import re
import json
from docx import Document
import hashlib

doc = Document("/Users/sachin/.hermes/attachments/7300+ GS English Medium.docx")

# Get all text
full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text.strip())
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                full_text.append(cell.text.strip())

text = "\n".join(full_text)
print(f"Total text length: {len(text)}")

# The format is:
# Question text (may span multiple lines)
# (A) option1    (B) option2    (C) option3    (D) option4  -- tab or space separated
# Ans.(X) (Exam)
# Exp: explanation

# Let's use a more robust approach: find all "Ans." patterns and work backwards
ans_pattern = r'Ans\.\s*\(([A-D])\)\s*(\([^)]+\))'
ans_matches = list(re.finditer(ans_pattern, text))
print(f"Answer matches: {len(ans_matches)}")

chapter_keywords = {
    'ancient history': 'Ancient History',
    'medieval history': 'Medieval History',
    'modern history': 'Modern History',
    'indian geography': 'Indian Geography',
    'world geography': 'World Geography',
    'indian polity': 'Indian Polity',
    'indian economy': 'Indian Economy',
    'physics': 'Physics',
    'chemistry': 'Chemistry',
    'biology': 'Biology',
    'computer': 'Computer',
    'miscellaneous': 'Miscellaneous',
    'general science': 'General Science',
    'current affairs': 'Current Affairs',
    'static gk': 'Static GK',
    'history': 'History',
    'geography': 'Geography',
    'polity': 'Polity',
    'economy': 'Economy',
}

def detect_chapter(text_block):
    lower = text_block.lower()
    for keyword, chapter in chapter_keywords.items():
        if keyword in lower:
            return chapter
    return 'General'

def clean_text(t):
    return re.sub(r'\s+', ' ', t).strip()

questions = []

for i, ans_match in enumerate(ans_matches):
    # Get the block from previous answer end to this answer start
    block_start = ans_matches[i-1].end() if i > 0 else 0
    block_end = ans_match.end()
    block = text[block_start:block_end]
    
    # Detect chapter
    current_chapter = detect_chapter(block)
    
    # Find options BEFORE the answer
    # Options are tab-separated: (A) text (B) text (C) text (D) text
    # They might be on 1 line or 2 lines
    opt_section_end = ans_match.start()
    opt_section_start = max(block_start, opt_section_end - 500)  # Look back 500 chars
    opt_section = text[opt_section_start:opt_section_end]
    
    # Find all (A) (B) (C) (D) patterns
    opt_pattern = r'\(([A-D])\)\s*([^\(]+?)(?=\s*\([A-D]\)|\s*Ans\.|\s*Exp:|\Z)'
    opt_matches = list(re.finditer(opt_pattern, opt_section, re.DOTALL))
    
    if len(opt_matches) >= 4:
        # Take last 4 options
        options = {}
        for m in opt_matches[-4:]:
            key = m.group(1)
            val = clean_text(m.group(2))
            # Clean up tab/newline artifacts
            val = re.sub(r'^[\(\s]+', '', val)
            val = re.sub(r'[\)\s]+$', '', val)
            options[key] = val
        
        # Question text is before the first of these 4 options
        first_opt_pos = opt_matches[-4].start()
        question_text = opt_section[:first_opt_pos].strip()
        
        # Clean question: remove any previous explanation
        question_text = re.sub(r'Exp:.*?(?=\n|\Z)', '', question_text, flags=re.DOTALL)
        question_text = re.sub(r'Note:.*?(?=\n|\Z)', '', question_text, flags=re.DOTALL)
        question_text = clean_text(question_text)
        
        # Skip if it's just an explanation
        if question_text.lower().startswith('exp') or question_text.lower().startswith('note'):
            continue
        
        if len(question_text) < 10:
            continue
        
        # Get answer and exam
        correct_answer = ans_match.group(1)
        exam_tag = ans_match.group(2).strip()
        
        # Get explanation AFTER this answer
        exp_start = ans_match.end()
        next_ans_start = ans_matches[i+1].start() if i+1 < len(ans_matches) else len(text)
        exp_block = text[exp_start:next_ans_start]
        
        exp_match = re.search(r'Exp:\s*(.*?)(?:\n\n|Ans\.|\Z)', exp_block, re.DOTALL)
        explanation = clean_text(exp_match.group(1)) if exp_match else ""
        
        q_hash = hashlib.sha256((exam_tag + question_text + str(options)).encode()).hexdigest()
        
        questions.append({
            'question': question_text[:2000],
            'options': [
                {'key': 'A', 'text': options.get('A', ''), 'isCorrect': correct_answer == 'A'},
                {'key': 'B', 'text': options.get('B', ''), 'isCorrect': correct_answer == 'B'},
                {'key': 'C', 'text': options.get('C', ''), 'isCorrect': correct_answer == 'C'},
                {'key': 'D', 'text': options.get('D', ''), 'isCorrect': correct_answer == 'D'},
            ],
            'correct_answer': correct_answer,
            'exam': exam_tag,
            'chapter': current_chapter,
            'explanation': explanation[:2000],
            'hash': q_hash
        })

print(f"\nTotal questions extracted: {len(questions)}")

# Stats
chapter_counts = {}
exam_counts = {}
for q in questions:
    chapter_counts[q['chapter']] = chapter_counts.get(q['chapter'], 0) + 1
    exam = q['exam']
    exam_counts[exam] = exam_counts.get(exam, 0) + 1

print("\nBy Chapter:")
for chap, count in sorted(chapter_counts.items(), key=lambda x: -x[1]):
    print(f"  {chap}: {count}")

print("\nBy Exam (top 20):")
for exam, count in sorted(exam_counts.items(), key=lambda x: -x[1])[:20]:
    print(f"  {exam}: {count}")

# Show sample
print("\n--- Sample Questions ---")
for q in questions[:5]:
    print(f"Q: {q['question'][:120]}")
    print(f"Chapter: {q['chapter']}, Exam: {q['exam']}, Answer: {q['correct_answer']}")
    print(f"Options: A={q['options'][0]['text'][:60]}, B={q['options'][1]['text'][:60]}")
    print(f"         C={q['options'][2]['text'][:60]}, D={q['options'][3]['text'][:60]}")
    print(f"Exp: {q['explanation'][:100]}")
    print()

# Save
with open('/tmp/gs_7300_questions_v4.json', 'w') as f:
    json.dump(questions, f, ensure_ascii=False, indent=2)

print(f"\nSaved to /tmp/gs_7300_questions_v4.json")