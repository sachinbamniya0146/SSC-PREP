"""
Parse 7300+ GS English Medium DOCX - Parse tables properly
"""
import re
import json
from docx import Document
import hashlib

doc = Document("/Users/sachin/.hermes/attachments/7300+ GS English Medium.docx")

# First, let's check tables
print(f"Number of tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    print(f"Table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    # Show first few rows
    for row_idx, row in enumerate(table.rows[:5]):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {row_idx}: {cells}")
    print()

# Also get paragraphs
para_count = sum(1 for p in doc.paragraphs if p.text.strip())
print(f"Non-empty paragraphs: {para_count}")