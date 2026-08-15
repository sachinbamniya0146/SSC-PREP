"""
Parse SSC Reasoning 7200 - Handle table cells and multi-line options
"""
import re
import json
from docx import Document

doc = Document("/Users/sachin/.hermes/attachments/ssc-reasoning-7200-tcs-mcq-chapter-wise-english-2nbsped_compress.docx")

# Extract ALL text including tables
all_text = []

for para in doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text.strip())

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                all_text.append(cell.text.strip())

text = "\n".join(all_text)
print(f"Total text length: {len(text)}")

# Better regex: find exam + question + 4 options anywhere in text
# Pattern: exam tag, then question text, then (a) opt (b) opt (c) opt (d) opt
# Options can be on same line or separate lines

# Split by exam tags
exam_pattern = r'(SSC (?:CPO|CGL|CHSL|MTS) .+? \((?:Morning|Afternoon|Evening)\))'
parts = re.split(exam_pattern, text)

questions = []
for i in range(1, len(parts), 2):
    exam = parts[i]
    content = parts[i+1] if i+1 < len(parts) else ""
    
    # Find all (a)...(b)...(c)...(d) sequences
    # Options can have line breaks
    opt_pattern = r'\(a\)\s*(.*?)\s*\(b\)\s*(.*?)\s*\(c\)\s*(.*?)\s*\(d\)\s*(.*?)(?=\(a\)|SSC (?:CPO|CGL|CHSL|MTS) |\Z)'
    
    matches = re.findall(opt_pattern, content, re.DOTALL | re.IGNORECASE)
    
    for match in matches:
        opt_a, opt_b, opt_c, opt_d = match
        # Clean
        opt_a = re.sub(r'\s+', ' ', opt_a).strip()
        opt_b = re.sub(r'\s+', ' ', opt_b).strip()
        opt_c = re.sub(r'\s+', ' ', opt_c).strip()
        opt_d = re.sub(r'\s+', ' ', opt_d).strip()
        
        # Skip if options look like garbage (too long = probably not options)
        if len(opt_a) > 200 or len(opt_b) > 200 or len(opt_c) > 200 or len(opt_d) > 200:
            continue
        
        # Find question text before this option set
        # Search backwards from this match position
        match_pos = content.find(f'(a){opt_a}')
        if match_pos == -1:
            match_pos = content.find(f'(a) {opt_a}')
        if match_pos == -1:
            continue
            
        question_text = content[:match_pos].strip()
        # Take last 500 chars as question
        question_text = question_text[-500:]
        
        questions.append({
            'exam': exam,
            'question': question_text,
            'options': {
                'A': opt_a,
                'B': opt_b,
                'C': opt_c,
                'D': opt_d,
            }
        })

print(f"Found {len(questions)} questions")

# Save
with open('/tmp/parsed_questions_v5.json', 'w') as f:
    json.dump(questions, f, indent=2, ensure_ascii=False)

# Stats
exam_counts = {}
for q in questions:
    exam = q['exam']
    exam_counts[exam] = exam_counts.get(exam, 0) + 1

print("\nExam distribution (top 30):")
for exam, count in sorted(exam_counts.items(), key=lambda x: -x[1])[:30]:
    print(f"  {exam}: {count}")

print(f"\nTotal: {sum(exam_counts.values())}")

# Show first 5
print("\nFirst 5 questions:")
for q in questions[:5]:
    print(f"\nExam: {q['exam']}")
    print(f"Q: {q['question'][:300]}...")
    print(f"Options: {q['options']}")

