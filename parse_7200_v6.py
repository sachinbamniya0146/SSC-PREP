"""
Parse SSC Reasoning 7200 - Handle options with better boundaries
"""
import re
import json
from docx import Document

doc = Document("/Users/sachin/.hermes/attachments/ssc-reasoning-7200-tcs-mcq-chapter-wise-english-2nbsped_compress.docx")

# Extract ALL text including tables
all_text = []
for para in doc.paragraphs:
    if para.text.strip():
        all_text.append(para.text.strip())

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                all_text.append(cell.text.strip())

text = "\n".join(all_text)

# Better approach: scan for complete question blocks
# A question block = exam tag + question + 4 options ending before next exam tag or question

# First, split by exam tags
exam_pattern = r'(SSC (?:CPO|CGL|CHSL|MTS) .+? \((?:Morning|Afternoon|Evening)\))'
parts = re.split(exam_pattern, text)

questions = []
for i in range(1, len(parts), 2):
    exam = parts[i]
    content = parts[i+1] if i+1 < len(parts) else ""
    
    # Find all question-option sets within this exam content
    # Pattern: question text followed by (a) opt (b) opt (c) opt (d) opt
    # Options end when next question starts (has its own (a)) or exam ends
    
    # Split content by (a) to find option starts
    opt_starts = [(m.start(), m.group()) for m in re.finditer(r'\(a\)', content)]
    
    for j, (start_pos, _) in enumerate(opt_starts):
        # Find the end of this option set (next (a) or end of content)
        end_pos = opt_starts[j+1][0] if j+1 < len(opt_starts) else len(content)
        
        option_block = content[start_pos:end_pos]
        
        # Parse options from this block
        opt_match = re.search(
            r'\(a\)\s*(.*?)\s*\(b\)\s*(.*?)\s*\(c\)\s*(.*?)\s*\(d\)\s*(.*?)(?=\s*(?:\([a-d]\)|SSC |\Z))',
            option_block, re.DOTALL | re.IGNORECASE
        )
        
        if not opt_match:
            continue
            
        opt_a = re.sub(r'\s+', ' ', opt_match.group(1)).strip()
        opt_b = re.sub(r'\s+', ' ', opt_match.group(2)).strip()
        opt_c = re.sub(r'\s+', ' ', opt_match.group(3)).strip()
        opt_d = re.sub(r'\s+', ' ', opt_match.group(4)).strip()
        
        # Skip if options too long (likely captured next question)
        if len(opt_a) > 150 or len(opt_b) > 150 or len(opt_c) > 150 or len(opt_d) > 150:
            continue
            
        # Get question text (text before this option block)
        # Find the previous question boundary
        question_end = start_pos
        question_text = content[:question_end].strip()
        
        # Clean up question text - take last reasonable portion
        # Remove any trailing incomplete options
        question_text = re.sub(r'\([a-d]\)\s*.*$', '', question_text, flags=re.DOTALL | re.IGNORECASE)
        question_text = question_text[-500:].strip()
        
        # Clean options: remove trailing question text that bled in
        for opt_key, opt_val in [('A', opt_a), ('B', opt_b), ('C', opt_c), ('D', opt_d)]:
            # Remove "Select the..." or "SSC " patterns that bled in
            opt_val = re.sub(r'\s+Select the.*$', '', opt_val, flags=re.IGNORECASE)
            opt_val = re.sub(r'\s+SSC (?:CPO|CGL|CHSL|MTS).*$', '', opt_val, flags=re.IGNORECASE)
            opt_val = re.sub(r'\s+In the following.*$', '', opt_val, flags=re.IGNORECASE)
            opt_val = re.sub(r'\s+Which.*\?.*$', '', opt_val, flags=re.IGNORECASE)
            opt_val = re.sub(r'\s+Identify.*$', '', opt_val, flags=re.IGNORECASE)
            if opt_key == 'A': opt_a = opt_val
            elif opt_key == 'B': opt_b = opt_val
            elif opt_key == 'C': opt_c = opt_val
            elif opt_key == 'D': opt_d = opt_val
        
        questions.append({
            'exam': exam,
            'question': question_text,
            'options': {'A': opt_a, 'B': opt_b, 'C': opt_c, 'D': opt_d}
        })

print(f"Found {len(questions)} questions")

# Save
with open('/tmp/parsed_questions_v6.json', 'w') as f:
    json.dump(questions, f, indent=2, ensure_ascii=False)

# Stats
exam_counts = {}
for q in questions:
    exam = q['exam']
    exam_counts[exam] = exam_counts.get(exam, 0) + 1

print("\nExam distribution (top 30):")
for exam, count in sorted(exam_counts.items(), key=lambda x: -x[1])[:30]:
    print(f"  {exam}: {count}")

print(f"\nTotal: {sum(exam_counts.values())}")

# Show first 10
print("\nFirst 10 questions:")
for q in questions[:10]:
    print(f"\nExam: {q['exam']}")
    print(f"Q: {q['question'][:300]}...")
    print(f"Options: {q['options']}")

