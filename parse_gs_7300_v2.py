"""
Parse 7300+ GS English Medium DOCX - Better parser
"""
import re
import json
from docx import Document
import hashlib

doc = Document("/Users/sachin/.hermes/attachments/7300+ GS English Medium.docx")

# Get all text
full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text.strip())
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                full_text.append(cell.text.strip())

text = "\n".join(full_text)
print(f"Total text length: {len(text)}")

# The format is:
# Question text
# Option A    Option B    Option C    Option D
# Ans.(X) (Exam)
# Exp: explanation

# Split by "Ans." to get question blocks
# But answers are inline, so let's use a different approach

# Find all answer positions
ans_pattern = r'Ans\.\s*\(([A-D])\)\s*(\([^)]+\))'
ans_matches = list(re.finditer(ans_pattern, text))
print(f"Answer matches: {len(ans_matches)}")

# For each answer, extract the question before it
questions = []
chapter_keywords = {
    'ancient history': 'Ancient History',
    'medieval history': 'Medieval History',
    'modern history': 'Modern History',
    'indian geography': 'Indian Geography',
    'world geography': 'World Geography',
    'indian polity': 'Indian Polity',
    'indian economy': 'Indian Economy',
    'physics': 'Physics',
    'chemistry': 'Chemistry',
    'biology': 'Biology',
    'computer': 'Computer',
    'miscellaneous': 'Miscellaneous',
    'general science': 'General Science',
    'current affairs': 'Current Affairs',
    'static gk': 'Static GK',
}

def detect_chapter(text_block):
    lower = text_block.lower()
    for keyword, chapter in chapter_keywords.items():
        if keyword in lower:
            return chapter
    return 'General'

def clean_text(t):
    return re.sub(r'\s+', ' ', t).strip()

# Track chapter
current_chapter = 'General'
last_ans_end = 0

for i, ans_match in enumerate(ans_matches):
    # Get text from last answer end to this answer start
    block_start = last_ans_end if i > 0 else 0
    block_end = ans_match.start()
    block = text[block_start:block_end]
    last_ans_end = ans_match.end()
    
    # Check for chapter in this block
    detected = detect_chapter(block)
    if detected != 'General':
        current_chapter = detected
    
    # Find options in block
    # Options are separated by tabs or newlines, format: (A) text (B) text (C) text (D) text
    opt_pattern = r'\(([A-D])\)\s*([^\n\(]+?)(?=\s*\([A-D]\)|Ans\.|Exp:|\Z)'
    opt_matches = list(re.finditer(opt_pattern, block, re.DOTALL))
    
    if len(opt_matches) >= 4:
        options = {}
        for m in opt_matches[:4]:
            options[m.group(1)] = clean_text(m.group(2))
        
        # Question text is before first option
        first_opt_pos = opt_matches[0].start()
        question_text = block[:first_opt_pos].strip()
        
        # Clean question text - remove any trailing exam tags
        question_text = re.sub(r'\([^)]*SSC [^)]*\)', '', question_text)
        question_text = re.sub(r'SSC [A-Z]+ \d{4}.*', '', question_text)
        question_text = clean_text(question_text)
        
        # Get answer and exam
        correct_answer = ans_match.group(1)
        exam_tag = ans_match.group(2).strip()
        
        # Get explanation (after this answer)
        exp_start = ans_match.end()
        exp_end = text.find('\n\n', exp_start)
        if exp_end == -1:
            exp_end = exp_start + 500
        exp_block = text[exp_start:exp_end]
        
        exp_match = re.search(r'Exp:\s*(.*?)(?:\n\n|\Z)', exp_block, re.DOTALL)
        explanation = clean_text(exp_match.group(1)) if exp_match else ""
        
        # Skip if question too short
        if len(question_text) < 5:
            question_text = "Question from " + exam_tag
        
        q_hash = hashlib.sha256((exam_tag + question_text + str(options)).encode()).hexdigest()
        
        questions.append({
            'question': question_text[:2000],
            'options': [
                {'key': 'A', 'text': options.get('A', ''), 'isCorrect': correct_answer == 'A'},
                {'key': 'B', 'text': options.get('B', ''), 'isCorrect': correct_answer == 'B'},
                {'key': 'C', 'text': options.get('C', ''), 'isCorrect': correct_answer == 'C'},
                {'key': 'D', 'text': options.get('D', ''), 'isCorrect': correct_answer == 'D'},
            ],
            'correct_answer': correct_answer,
            'exam': exam_tag,
            'chapter': current_chapter,
            'explanation': explanation[:2000],
            'hash': q_hash
        })

print(f"\nTotal questions extracted: {len(questions)}")

# Stats
chapter_counts = {}
exam_counts = {}
for q in questions:
    chapter_counts[q['chapter']] = chapter_counts.get(q['chapter'], 0) + 1
    exam = q['exam']
    exam_counts[exam] = exam_counts.get(exam, 0) + 1

print("\nBy Chapter:")
for chap, count in sorted(chapter_counts.items(), key=lambda x: -x[1]):
    print(f"  {chap}: {count}")

print("\nBy Exam (top 20):")
for exam, count in sorted(exam_counts.items(), key=lambda x: -x[1])[:20]:
    print(f"  {exam}: {count}")

# Show sample
print("\n--- Sample Questions ---")
for q in questions[:3]:
    print(f"Q: {q['question'][:100]}")
    print(f"Chapter: {q['chapter']}, Exam: {q['exam']}, Answer: {q['correct_answer']}")
    print(f"Exp: {q['explanation'][:100]}")
    print()

# Save
with open('/tmp/gs_7300_questions_v2.json', 'w') as f:
    json.dump(questions, f, ensure_ascii=False, indent=2)

print(f"\nSaved to /tmp/gs_7300_questions_v2.json")