"""
Parse 7300+ GS English Medium DOCX - Extract questions with answers & explanations
"""
import re
import json
from docx import Document
import hashlib

doc = Document("/Users/sachin/.hermes/attachments/7300+ GS English Medium.docx")

# Get all text
full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text.strip())
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                full_text.append(cell.text.strip())

text = "\n".join(full_text)
print(f"Total text length: {len(text)}")

# Pattern to find questions with options, answers, and explanations
# Format: Question text -> Options (A) (B) (C) (D) -> Ans. (X) (Exam) -> Exp: explanation

# First, let's find all answer patterns
ans_pattern = r'Ans\.\s*\(([A-D])\)\s*(\([^)]+\))'
ans_matches = list(re.finditer(ans_pattern, text))
print(f"Answer patterns found: {len(ans_matches)}")

# Find all explanation patterns
exp_pattern = r'Exp:\s*(.*?)(?=\n\n|\t\t|\Z)'
exp_matches = list(re.finditer(exp_pattern, text, re.DOTALL))
print(f"Explanation patterns found: {len(exp_matches)}")

# Better approach: Split by double newlines and process blocks
blocks = re.split(r'\n\s*\n', text)

questions = []
current_question = None
current_options = {}
current_exam = ""
current_chapter = "General"

# Chapter detection
chapter_keywords = {
    'ancient history': 'Ancient History',
    'medieval history': 'Medieval History',
    'modern history': 'Modern History',
    'indian geography': 'Indian Geography',
    'world geography': 'World Geography',
    'indian polity': 'Indian Polity',
    'indian economy': 'Indian Economy',
    'physics': 'Physics',
    'chemistry': 'Chemistry',
    'biology': 'Biology',
    'computer': 'Computer',
    'miscellaneous': 'Miscellaneous',
    'general science': 'General Science',
    'current affairs': 'Current Affairs',
    'static gk': 'Static GK',
}

def detect_chapter(text_block):
    lower = text_block.lower()
    for keyword, chapter in chapter_keywords.items():
        if keyword in lower:
            return chapter
    return None

def clean_text(t):
    return re.sub(r'\s+', ' ', t).strip()

for block in blocks:
    block = block.strip()
    if not block:
        continue
    
    # Check for chapter headers
    chapter = detect_chapter(block)
    if chapter:
        current_chapter = chapter
        continue
    
    # Check for answer pattern
    ans_match = re.search(ans_pattern, block)
    if ans_match:
        correct_answer = ans_match.group(1)
        exam_tag = ans_match.group(2).strip()
        current_exam = exam_tag
        
        # Extract explanation if present
        exp_match = re.search(r'Exp:\s*(.*?)(?:\n|$)', block)
        explanation = exp_match.group(1).strip() if exp_match else ""
        
        # If we have a pending question, complete it
        if current_question and current_options:
            # Build options JSON
            options_list = []
            for key in ['A', 'B', 'C', 'D']:
                if key in current_options:
                    options_list.append({
                        'key': key,
                        'text': clean_text(current_options[key]),
                        'isCorrect': (key == correct_answer)
                    })
            
            if len(options_list) == 4:
                q_hash = hashlib.sha256((current_exam + current_question + str(current_options)).encode()).hexdigest()
                
                questions.append({
                    'question': clean_text(current_question),
                    'options': options_list,
                    'correct_answer': correct_answer,
                    'exam': current_exam,
                    'chapter': current_chapter,
                    'explanation': clean_text(explanation),
                    'hash': q_hash
                })
        
        # Reset for next question
        current_question = None
        current_options = {}
        continue
    
    # Check for option patterns
    opt_matches = re.findall(r'\(([A-D])\)\s*([^\n\(]+)', block)
    if opt_matches:
        for key, val in opt_matches:
            current_options[key] = val.strip()
        
        # The question text is before the first option
        first_opt_pos = block.find(opt_matches[0][1])
        if first_opt_pos > 0:
            current_question = block[:first_opt_pos].strip()

# Stats
print(f"\nTotal questions extracted: {len(questions)}")

chapter_counts = {}
exam_counts = {}
for q in questions:
    chapter_counts[q['chapter']] = chapter_counts.get(q['chapter'], 0) + 1
    exam = q['exam']
    exam_counts[exam] = exam_counts.get(exam, 0) + 1

print("\nBy Chapter:")
for chap, count in sorted(chapter_counts.items(), key=lambda x: -x[1]):
    print(f"  {chap}: {count}")

print("\nBy Exam (top 20):")
for exam, count in sorted(exam_counts.items(), key=lambda x: -x[1])[:20]:
    print(f"  {exam}: {count}")

# Save
with open('/tmp/gs_7300_questions.json', 'w') as f:
    json.dump(questions, f, ensure_ascii=False, indent=2)

print(f"\nSaved to /tmp/gs_7300_questions.json")