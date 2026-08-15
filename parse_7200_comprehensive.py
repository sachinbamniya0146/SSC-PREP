"""
Comprehensive SSC Reasoning 7200 Parser - Better version with chapter/exam tracking
"""
import re
import json
from docx import Document
import hashlib

doc = Document("/Users/sachin/.hermes/attachments/ssc-reasoning-7200-tcs-mcq-chapter-wise-english-2nbsped_compress (1).docx")

# Get all text
full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text.strip())
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                full_text.append(cell.text.strip())

text = "\n".join(full_text)
print(f"Total text length: {len(text)}")

# ============================================
# STEP 1: Parse chapter index (lines 1-555)
# ============================================
chapter_pattern = r'(\d+)\.\s*([A-Z][A-Z\s]+)\s*(\d+)\s*(\d+-\d+)'
chapters = {}
chapter_matches = re.finditer(chapter_pattern, text[:20000])
for m in chapter_matches:
    num = int(m.group(1))
    name = m.group(2).strip()
    count = int(m.group(3))
    pages = m.group(4)
    chapters[num] = {'name': name, 'count': count, 'pages': pages}
    print(f"Chapter {num}: {name} - {count} questions ({pages})")

print(f"\nTotal chapters: {len(chapters)}")
total_chapter_qs = sum(c['count'] for c in chapters.values())
print(f"Total from chapters: {total_chapter_qs}")

# ============================================
# STEP 2: Parse exam-wise question counts
# ============================================
exam_section = text[10000:30000]
exam_pattern = r'(\d+)\.\s*(SSC [A-Z]+ \d{4})\s*(\d+)\s*(\d+)\s*(\d+)'
exam_matches = re.finditer(exam_pattern, exam_section)
exams = {}
for m in exam_matches:
    exams[m.group(2)] = {'per_shift': int(m.group(3)), 'shifts': int(m.group(4)), 'total': int(m.group(5))}
    print(f"Exam: {m.group(2)} - {m.group(3)} q/shift x {m.group(4)} shifts = {m.group(5)}")

# ============================================
# STEP 3: Parse ALL questions using better regex
# ============================================
# Find all question blocks - they start with exam tag and end before next exam tag
exam_tag_pattern = r'(SSC (?:CPO|CGL|CHSL|MTS|CPO|CHSL|CGL|MTS) .+? \((?:Morning|Afternoon|Evening)\))'

# Split text by exam tags
parts = re.split(f'({exam_tag_pattern})', text)

questions = []
current_exam = "Unknown"
current_chapter = "Unknown"
chapter_names_list = list(chapters.values())

# Track chapter based on question count progress
chapter_bounds = {}
cum = 0
for i, (num, info) in enumerate(chapters.items()):
    chapter_bounds[i] = (cum, cum + info['count'], info['name'])
    cum += info['count']

question_num = 0

# Better approach: find all (a)...(b)...(c)...(d) patterns with context
# Pattern: exam tag -> question text -> (a) optA (b) optB (c) optC (d) optD
opt_block_pattern = r'\(a\)\s*(.*?)\s*\(b\)\s*(.*?)\s*\(c\)\s*(.*?)\s*\(d\)\s*(.*?)(?=\s*\(a\)|SSC (?:CPO|CGL|CHSL|MTS) |\Z)'

opt_matches = list(re.finditer(opt_block_pattern, text, re.DOTALL | re.IGNORECASE))
print(f"\nTotal option blocks found: {len(opt_matches)}")

for i, match in enumerate(opt_matches):
    opt_a = re.sub(r'\s+', ' ', match.group(1)).strip()
    opt_b = re.sub(r'\s+', ' ', match.group(2)).strip()
    opt_c = re.sub(r'\s+', ' ', match.group(3)).strip()
    opt_d = re.sub(r'\s+', ' ', match.group(4)).strip()
    
    # Skip if options too long (likely not real options)
    if max(len(opt_a), len(opt_b), len(opt_c), len(opt_d)) > 200:
        continue
    
    # Get text before this match to find exam tag and question
    before = text[:match.start()]
    
    # Find latest exam tag
    exam_matches_before = list(re.finditer(exam_tag_pattern, before))
    if exam_matches_before:
        current_exam = exam_matches_before[-1].group(1)
    
    # Find question text (between previous options end and this options start)
    prev_opt_end = 0
    for m in re.finditer(r'\(d\)\s*.*?(?=\s*\(a\)|SSC )', before):
        prev_opt_end = m.end()
    
    if not prev_opt_end and exam_matches_before:
        prev_opt_end = exam_matches_before[-1].end()
    
    question_text = before[prev_opt_end:match.start()].strip()
    question_text = re.sub(r'\s+', ' ', question_text).strip()
    
    # Clean up question text
    question_text = re.sub(r'^.*?(?:Select the|In the following|Which of|Find the|Identify)\s*', '', question_text, flags=re.IGNORECASE)
    question_text = re.sub(r'SSC (?:CPO|CGL|CHSL|MTS) .+?\((?:Morning|Afternoon|Evening)\)', '', question_text)
    question_text = question_text.strip()
    
    # Skip if too short
    if len(question_text) < 5:
        # Try to infer from context
        question_text = f"Question from {current_exam}"
    
    # Clean options - remove bleed-in text
    def clean_opt(opt):
        opt = re.sub(r'\s+Select the.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+SSC (?:CPO|CGL|CHSL|MTS).*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+In the following.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+Which.*\?.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+Identify.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+Study the.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+Find the.*$', '', opt, flags=re.IGNORECASE)
        return opt.strip()
    
    opt_a = clean_opt(opt_a)
    opt_b = clean_opt(opt_b)
    opt_c = clean_opt(opt_c)
    opt_d = clean_opt(opt_d)
    
    # Determine chapter based on question number
    chapter_name = "General"
    for idx, (start, end, name) in chapter_bounds.items():
        if question_num >= start and question_num < end:
            chapter_name = name
            break
    
    # Extract year and shift from exam tag
    year_match = re.search(r'20\d{2}', current_exam)
    year = int(year_match.group()) if year_match else 2025
    
    shift = 'Morning' if '(Morning)' in current_exam else ('Afternoon' if '(Afternoon)' in current_exam else ('Evening' if '(Evening)' in current_exam else 'All'))
    
    # Determine exam ID
    exam_id_map = {
        'SSC CPO': 'exam-cpo',
        'SSC CGL': 'exam-cgl', 
        'SSC CHSL': 'exam-chsl',
        'SSC MTS': 'exam-mts',
    }
    exam_id = 'exam-cpo'
    for k, v in exam_id_map.items():
        if k in current_exam:
            exam_id = v
            break
    
    # Hash for deduplication
    hash_val = hashlib.sha256(question_text.lower().encode()).hexdigest()
    
    questions.append({
        'num': question_num + 1,
        'exam': current_exam,
        'exam_id': exam_id,
        'chapter': chapter_name,
        'year': year,
        'shift': shift,
        'question': question_text[:2000],
        'options': {'A': opt_a, 'B': opt_b, 'C': opt_c, 'D': opt_d},
        'hash': hash_val
    })
    
    question_num += 1

print(f"\nTotal clean questions extracted: {len(questions)}")

# Stats by exam
exam_counts = {}
chapter_counts = {}
for q in questions:
    exam_counts[q['exam']] = exam_counts.get(q['exam'], 0) + 1
    chapter_counts[q['chapter']] = chapter_counts.get(q['chapter'], 0) + 1

print("\nBy Exam:")
for exam, count in sorted(exam_counts.items(), key=lambda x: -x[1]):
    print(f"  {exam}: {count}")

print("\nBy Chapter (top 15):")
for chap, count in sorted(chapter_counts.items(), key=lambda x: -x[1])[:15]:
    print(f"  {chap}: {count}")

# Save
with open('/tmp/questions_7200_v7.json', 'w') as f:
    json.dump(questions, f, ensure_ascii=False, indent=2)

print(f"\nSaved to /tmp/questions_7200_v7.json")