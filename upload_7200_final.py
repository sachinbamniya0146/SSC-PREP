"""
Upload 7200 questions to database - improved version
"""
import json
import re
import hashlib
import sys

# Load parsed questions
with open('/tmp/parsed_questions_v6.json') as f:
    questions = json.load(f)

# Read the full text again for better question extraction
from docx import Document
doc = Document("/Users/sachin/.hermes/attachments/ssc-reasoning-7200-tcs-mcq-chapter-wise-english-2nbsped_compress.docx")

full_text = []
for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text.strip())
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                full_text.append(cell.text.strip())

text = "\n".join(full_text)

# Better extraction: find each (a)...(b)...(c)...(d) and get question before it
questions_clean = []

# Find all option blocks with their positions
opt_pattern = r'\(a\)\s*(.*?)\s*\(b\)\s*(.*?)\s*\(c\)\s*(.*?)\s*\(d\)\s*(.*?)(?=\s*\(a\)|SSC (?:CPO|CGL|CHSL|MTS) |\Z)'
matches = list(re.finditer(opt_pattern, text, re.DOTALL | re.IGNORECASE))

print(f"Total option blocks found: {len(matches)}")

for match in matches:
    opt_a = re.sub(r'\s+', ' ', match.group(1)).strip()
    opt_b = re.sub(r'\s+', ' ', match.group(2)).strip()
    opt_c = re.sub(r'\s+', ' ', match.group(3)).strip()
    opt_d = re.sub(r'\s+', ' ', match.group(4)).strip()
    
    # Skip if too long
    if max(len(opt_a), len(opt_b), len(opt_c), len(opt_d)) > 150:
        continue
    
    # Get exam tag before this match
    before_text = text[:match.start()]
    exam_match = re.search(r'(SSC (?:CPO|CGL|CHSL|MTS) .+? \((?:Morning|Afternoon|Evening)\))\s*$', before_text)
    exam = exam_match.group(1) if exam_match else "Unknown"
    
    # Get question text (text between previous exam/question and this options)
    # Find previous option block end or exam tag
    prev_end = 0
    for m in re.finditer(r'\(d\)\s*.*?(?=\s*\(a\)|SSC )', before_text):
        prev_end = m.end()
    
    if not prev_end:
        exam_tag_match = re.search(r'(SSC (?:CPO|CGL|CHSL|MTS) .+? \((?:Morning|Afternoon|Evening)\))', before_text)
        if exam_tag_match:
            prev_end = exam_tag_match.end()
    
    question_text = before_text[prev_end:match.start()].strip()
    
    # Clean question
    question_text = re.sub(r'\s+', ' ', question_text).strip()
    question_text = re.sub(r'^.*?Questions?\s*[:-]', '', question_text, flags=re.IGNORECASE)
    question_text = re.sub(r'SSC (?:CPO|CGL|CHSL|MTS) .+?\((?:Morning|Afternoon|Evening)\)', '', question_text)
    question_text = question_text.strip()
    
    # Skip if question too short or looks like garbage
    if len(question_text) < 10:
        continue
    
    # Clean options - remove bleed-in
    for opt in [opt_a, opt_b, opt_c, opt_d]:
        opt = re.sub(r'\s+Select the.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+SSC (?:CPO|CGL|CHSL|MTS).*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+In the following.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+Which.*\?.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+Identify.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+Study the.*$', '', opt, flags=re.IGNORECASE)
        opt = re.sub(r'\s+Find the.*$', '', opt, flags=re.IGNORECASE)
    
    # Extract year/shift
    year_match = re.search(r'20\d{2}', exam)
    year = int(year_match.group()) if year_match else 2025
    
    shift = 'Morning' if '(Morning)' in exam else ('Afternoon' if '(Afternoon)' in exam else ('Evening' if '(Evening)' in exam else 'All'))
    
    # Determine exam ID
    exam_id_map = {
        'SSC CPO': 'exam-cpo',
        'SSC CGL': 'exam-cgl', 
        'SSC CHSL': 'exam-chsl',
        'SSC MTS': 'exam-mts',
    }
    exam_id = 'exam-cpo'
    for k, v in exam_id_map.items():
        if k in exam:
            exam_id = v
            break
    
    # Hash
    hash_val = hashlib.sha256(question_text.lower().encode()).hexdigest()
    
    questions_clean.append({
        'exam': exam,
        'exam_id': exam_id,
        'year': year,
        'shift': shift,
        'question': question_text[:2000],
        'options': {'A': opt_a, 'B': opt_b, 'C': opt_c, 'D': opt_d},
        'hash': hash_val
    })

print(f"Clean questions: {len(questions_clean)}")

# Write SQL file
with open('/tmp/upload_7200.sql', 'w') as f:
    f.write('BEGIN;\n')
    
    for q in questions_clean:
        opts = q['options']
        options_json = json.dumps([
            {'key': 'A', 'text': opts['A'].replace("'", "''"), 'isCorrect': False},
            {'key': 'B', 'text': opts['B'].replace("'", "''"), 'isCorrect': False},
            {'key': 'C', 'text': opts['C'].replace("'", "''"), 'isCorrect': False},
            {'key': 'D', 'text': opts['D'].replace("'", "''"), 'isCorrect': False},
        ])
        
        question_escaped = q['question'].replace("'", "''")
        
        sql = f"""INSERT INTO questions (id, "questionText", "questionTextHindi", "optionsJson", "correctAnswer", explanation, "explanationSource", "searchHash", "sourcePdfId", "importBatchId", "subjectId", "examId", year, shift, marks, "negativeMarks", "isApproved", "isActive", "answerVerificationStatus", "aiConfidenceScore", "reviewStatus", "createdAt", "updatedAt")
VALUES (gen_random_uuid(), '{question_escaped}', NULL, '{options_json}'::jsonb, 'A', NULL, 'PDF'::"ExplanationSource", '{q['hash']}', '8f434dff-f8e0-4055-ba82-908e3a429690', '46a7b2d3-9439-4f9c-aeac-cb042e12a8f7', 'c5d2bb1f-ac87-432d-9462-954835c4a4ed', '{q['exam_id']}', {q['year']}, '{q['shift']}', 1, 0.25, false, true, 'UNVERIFIED_SINGLE_SOURCE', 0.5, 'AI_DRAFT', NOW(), NOW())
ON CONFLICT ("searchHash") DO NOTHING;\n"""
        f.write(sql)
    
    f.write('COMMIT;\n')

print("SQL written to /tmp/upload_7200.sql")